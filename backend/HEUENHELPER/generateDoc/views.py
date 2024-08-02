import os

from django.http import FileResponse
from docx.oxml.ns import qn
# Create your views here.
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status
from rest_framework.permissions import IsAuthenticated
from docx import Document
from services.AI_API_CALL import RateLimitException, AIAdviceService, InvalidRoleException

class GenerateLessonPlanView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request, *args, **kwargs):
        prompt = request.data.get('prompt')
        user_id = request.user.id

        if not prompt:
            return Response({"error": "教案提示词不能为空"}, status=status.HTTP_400_BAD_REQUEST)

        # 调用AI生成建议
        ai_service = AIAdviceService()
        ai_service.add_param('system', '你现在是一位老师，请根据下面的提示词生成教案。教案应包括以下部分：1. 教案标题 2. 教学目标 3. 教学过程 4. 课后作业 用中文回答。这四个部分应该分别分行，并以1. 2. 3. 4. 分别开头。整个回答应该只有四行。要详细回答，设计教学步骤。').add_param('user', prompt)
        advice = ai_service.get_advice(user_id=user_id)
        advice = advice['choices'][0]['message']['content']
        # 检查返回内容是否符合预期格式
        def check_format(advice):
            required_sections = ["1. 教案标题", "2. 教学目标", "3. 教学过程", "4. 课后作业"]
            for section in required_sections:
                if section not in advice:
                    return False
            return True

        # 如果不符合格式，重新生成
        max_retries = 3
        retries = 0
        while not check_format(advice) and retries < max_retries:
            advice = ai_service.get_advice(user_id=user_id)
            retries += 1

        if not check_format(advice):
            return Response({"error": "AI生成的教案格式不正确，请重试。"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

        # 生成docx文件
        doc = Document()
        doc.add_heading('教案生成', 0)
        doc.styles['Normal'].font.name = 'Times New Roman'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        # 解析AI返回的建议并添加到doc中
        advice_lines = advice.split('\n')
        for line in advice_lines:
            if line.startswith("1. 教案标题"):
                r = doc.add_heading("", level=1)
                r = r.add_run(line.replace("1. 教案标题：", "").strip())
                r.font.name=u'Times New Roman'
                r._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            elif line.startswith("2. 教学目标"):
                r = doc.add_heading("", level=2)
                r = r.add_run("教学目标")
                doc.add_paragraph(line.replace("2. 教学目标：", "").strip())

                r.font.name = u'Times New Roman'
                r._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            elif line.startswith("3. 教学过程"):
                r = doc.add_heading("", level=2)
                r = r.add_run("教学过程")
                doc.add_paragraph(line.replace("3. 教学过程：", "").strip())
                r.font.name = u'Times New Roman'
                r._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            elif line.startswith("4. 课后作业"):
                r = doc.add_heading("", level=2)
                r = r.add_run("课后作业")
                doc.add_paragraph(line.replace("4. 课后作业：", "").strip())
                r.font.name = u'Times New Roman'
                r._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            else:
                doc.add_paragraph(line.strip())

        output_dir = './tmp'
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)

        doc_path = os.path.join(output_dir, f'lesson_plan_{user_id}.docx')
        doc.save(doc_path)

        # 读取生成的docx文件并返回
        response = FileResponse(open(doc_path, 'rb'),filename = f"lesson_plan_{user_id}.docx", content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', as_attachment=True, status=status.HTTP_200_OK)
        #response['Content-Disposition'] = f'attachment; filename="lesson_plan_{user_id}.docx"'
        return response

        return Response({"error": "未能生成教案"}, status=status.HTTP_400_BAD_REQUEST)
